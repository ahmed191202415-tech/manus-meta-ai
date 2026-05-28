from typing import Any, Dict

from docx import Document

from app.config import EXPORT_DIR
from app.utils.files import sanitize_file_name


def save_docx_report_local(body) -> Dict[str, Any]:
    file_name = sanitize_file_name(body.file_name, "report", "docx")
    output_path = EXPORT_DIR / file_name

    document = Document()
    document.add_heading(body.title, level=0)
    if body.subtitle:
        document.add_paragraph(body.subtitle)

    for section in body.sections:
        if section.heading:
            document.add_heading(section.heading, level=1)
        for paragraph in section.paragraphs:
            document.add_paragraph(str(paragraph))
        if section.table_headers:
            table = document.add_table(rows=1, cols=len(section.table_headers))
            table.style = "Table Grid"
            for index, header in enumerate(section.table_headers):
                table.rows[0].cells[index].text = str(header)
            for row in section.table_rows:
                cells = table.add_row().cells
                for index, value in enumerate(row[:len(section.table_headers)]):
                    cells[index].text = str(value)

    document.save(output_path)
    return {
        "success": True,
        "file_name": file_name,
        "file_path": str(output_path),
        "export_dir": str(EXPORT_DIR),
    }
